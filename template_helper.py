import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document

def read_docx(file_path):
    """
    Извлечение текста из DOCX (без внешних зависимостей).
    """
    if not os.path.exists(file_path):
        return f"Ошибка: Файл не найден по пути {file_path}"
    try:
        with zipfile.ZipFile(file_path) as z:
            xml_content = z.read('word/document.xml')
            root = ET.fromstring(xml_content)
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            paragraphs = []
            for p in root.findall('.//w:p', namespaces):
                texts = p.findall('.//w:t', namespaces)
                if texts:
                    p_text = "".join([t.text for t in texts if t.text])
                    paragraphs.append(p_text)
            return "\n".join(paragraphs)
    except Exception as e:
        return f"Ошибка при чтении DOCX: {e}"

def replace_placeholders(doc, data):
    """
    Безопасная замена плейсхолдеров в параграфах и таблицах документа Word.
    """
    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for key, val in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in p.text:
                    # Попытка заменить внутри конкретных прогонов (runs) для сохранения форматирования
                    replaced_in_runs = False
                    for run in p.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(val))
                            replaced_in_runs = True
                    # Если плейсхолдер был разбит на несколько runs, делаем грубую замену в тексте параграфа
                    if not replaced_in_runs or placeholder in p.text:
                        p.text = p.text.replace(placeholder, str(val))

    # Замена в основном тексте
    replace_in_paragraphs(doc.paragraphs)
    
    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

def fill_template(template_path, output_path, data):
    """
    Загружает шаблон, заполняет его данными и сохраняет результат.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Файл шаблона не найден: {template_path}")
    
    doc = Document(template_path)
    replace_placeholders(doc, data)
    doc.save(output_path)
    return output_path

def create_default_template(target_path):
    """
    Создает демонстрационный шаблон DOCX, если у пользователя нет готового.
    """
    doc = Document()
    doc.add_heading("ПРЕДСТАВЛЕНИЕ В ПРИКАЗ", level=1)
    
    doc.add_paragraph("Настоящим документом подтверждается направление студентов на учебную/производственную практику.")
    
    p1 = doc.add_paragraph()
    p1.add_run("Студент (ФИО): ").bold = True
    p1.add_run("{{student_fio}}\n")
    p1.add_run("Группа: ").bold = True
    p1.add_run("{{group_name}}\n")
    p1.add_run("Направление подготовки: ").bold = True
    p1.add_run("{{direction}}\n")
    p1.add_run("Кафедра: ").bold = True
    p1.add_run("{{department_name}}\n")
    
    doc.add_heading("Сведения о практике", level=2)
    p2 = doc.add_paragraph()
    p2.add_run("Тип практики: ").bold = True
    p2.add_run("{{practice_type}}\n")
    p2.add_run("Организация: ").bold = True
    p2.add_run("{{organization_name}}\n")
    p2.add_run("Приказ от: ").bold = True
    p2.add_run("{{order_date}} (Подписал: {{order_signed_by}})\n")
    
    doc.add_heading("Руководитель от кафедры", level=2)
    p3 = doc.add_paragraph()
    p3.add_run("Преподаватель (ФИО): ").bold = True
    p3.add_run("{{teacher_fio}}\n")
    p3.add_run("Должность: ").bold = True
    p3.add_run("{{teacher_position}} ({{teacher_degree}}, {{teacher_rank}})\n")
    p3.add_run("Контракт №: ").bold = True
    p3.add_run("{{contract_number}} от {{contract_date}}\n")
    
    # Добавим таблицу
    doc.add_heading("Академическая нагрузка по предмету", level=2)
    table = doc.add_table(rows=4, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр'
    hdr_cells[1].text = 'Значение'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Дисциплина'
    row_cells[1].text = '{{subject_name}}'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Часы лекций'
    row_cells[1].text = '{{subject_lectures}}'
    
    row_cells = table.rows[3].cells
    row_cells[0].text = 'Часы лабораторных'
    row_cells[1].text = '{{subject_labs}}'
    
    doc.save(target_path)
    print(f"Создан демонстрационный шаблон: {target_path}")
